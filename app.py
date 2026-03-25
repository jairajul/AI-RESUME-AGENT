import streamlit as st
import os
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 🔑 Read API key from Streamlit Secrets
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Debug line (temporary)
st.write("API key loaded:", os.getenv("OPENAI_API_KEY") is not None)


# 🧾 Create clean Word Resume
def create_word_doc(text):
    doc = Document()

    # Keep only resume part
    if "=== ANALYSIS ===" in text:
        text = text.split("=== ANALYSIS ===")[0]

    text = text.replace("=== RESUME ===", "")
    text = text.replace("---", "").replace("___", "")

    lines = text.split("\n")

    for i, line in enumerate(lines):
        line = line.strip()

        if not line:
            continue

        # Name (First line)
        if i == 0:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section headings
        elif line.isupper():
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)

        # Bullet points
        elif line.startswith("-"):
            doc.add_paragraph(line.replace("-", "").strip(), style="List Bullet")

        else:
            doc.add_paragraph(line)

    file_path = "Rajul_Tailored_Resume.docx"
    doc.save(file_path)

    return file_path


# 🎯 UI
st.title("🤖 AI Resume + Job Agent")

jd = st.text_area("📄 Paste Job Description")
resume = st.text_area("📄 Paste Your Resume")

if st.button("🚀 Analyze"):

    if jd and resume:

        prompt = f"""
You are an expert hiring manager and career strategist.

STRICT FORMAT:

=== RESUME ===
Write ONLY the tailored resume here.
Clean, professional, no explanations.

=== ANALYSIS ===
Include:
- Match Score
- Strengths
- Gaps
- Keywords
- Interview Questions & Answers
- Action Plan
- Talking Points

Candidate Profile:
{resume}

Job Description:
{jd}
"""

        with st.spinner("🤖 Analyzing..."):
            response = client.chat.completions.create(
                model="gpt-4.1",
                messages=[{"role": "user", "content": prompt}]
            )

            output = response.choices[0].message.content

        # 🔀 Split Resume & Analysis
        resume_part = ""
        analysis_part = ""

        if "=== RESUME ===" in output and "=== ANALYSIS ===" in output:
            resume_part = output.split("=== RESUME ===")[1].split("=== ANALYSIS ===")[0]
            analysis_part = output.split("=== ANALYSIS ===")[1]
        else:
            analysis_part = output

        # 📊 Display ANALYSIS only
        st.subheader("📊 Analysis")

        sections = analysis_part.split("\n\n")

        for section in sections:
            if "Match Score" in section:
                st.success(section)

            elif "Strengths" in section:
                st.info(section)

            elif "Gaps" in section:
                st.warning(section)

            elif "Action Plan" in section:
                st.success(section)

            elif "Interview" in section:
                st.info(section)

            else:
                st.markdown(section)

        # 📥 Download ONLY Resume
        if resume_part:
            file_path = create_word_doc(resume_part)

            with open(file_path, "rb") as file:
                st.download_button(
                    label="📥 Download Final Resume",
                    data=file,
                    file_name="Rajul_Tailored_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    else:
        st.warning("⚠️ Please fill both fields!")
